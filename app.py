from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime

app = Flask(__name__)
UPLOAD_FOLDER = 'static/uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def substituir_placeholders(paragraphs, dados):
    for p in paragraphs:
        for chave, valor in dados.items():
            alvo = f"{{{{{chave}}}}}"
            if alvo in p.text:
                for run in p.runs:
                    if alvo in run.text:
                        run.text = run.text.replace(alvo, valor)

def preencher_docx(dados, fotos):
    doc = Document("MODELO_LAUDO.docx")

    # 1) Corpo do documento
    substituir_placeholders(doc.paragraphs, dados)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                substituir_placeholders(cell.paragraphs, dados)

    # 2) Cabeçalhos e rodapés
    for section in doc.sections:
        substituir_placeholders(section.header.paragraphs, dados)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    substituir_placeholders(cell.paragraphs, dados)
        substituir_placeholders(section.footer.paragraphs, dados)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    substituir_placeholders(cell.paragraphs, dados)

    # 3) Inserção das fotos na última tabela do corpo
    tabelas = doc.tables
    if tabelas:
        tabela = tabelas[-1]
        for i in range(1, 5):
            antes = fotos.get(f"foto{i}_antes")
            depois = fotos.get(f"foto{i}_depois")
            if antes:
                tabela.cell(i, 0).paragraphs[0].add_run().add_picture(antes, width=Inches(2))
            if depois:
                tabela.cell(i, 1).paragraphs[0].add_run().add_picture(depois, width=Inches(2))

    # 4) Salvar com nome único
    nome_arquivo = f"laudo_{dados['nroOS']}_{datetime.now():%Y%m%d%H%M%S}.docx"
    caminho = os.path.join(app.config['UPLOAD_FOLDER'], nome_arquivo)
    doc.save(caminho)
    return caminho

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        dados = {
            "nroOS": request.form.get("nroOS", ""),
            "DATA": request.form.get("DATA", ""),
            "Nome_Tecnico": request.form.get("Nome_Tecnico", ""),
            "Nome_Cliente": request.form.get("Nome_Cliente", ""),
            "Endereco_Cliente": request.form.get("Endereco_Cliente", ""),
            "Telefone_cliente": request.form.get("Telefone_cliente", ""),
            "Modelo_equipamento": request.form.get("Modelo_equipamento", ""),
            "Numero_Serie": request.form.get("Numero_Serie", ""),
            "Chamado_Aberto": request.form.get("Chamado_Aberto", ""),
            "Defeitos_Encontrados": request.form.get("Defeitos_Encontrados", ""),
            "Tarefas_Executadas": request.form.get("Tarefas_Executadas", ""),
        }

        fotos = {}
        for i in range(1, 5):
            for tipo in ("antes", "depois"):
                file = request.files.get(f"foto{i}_{tipo}")
                if file and file.filename:
                    path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
                    file.save(path)
                    fotos[f"foto{i}_{tipo}"] = path

        saída = preencher_docx(dados, fotos)
        return send_file(saída, as_attachment=True)

    return render_template("formulario.html")

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
